#!/usr/bin/env python3
"""Convert REPORT.md to REPORT.docx"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def parse_markdown_file(filepath):
    """Parse markdown file and return content structure"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

def add_text_to_doc(doc, text):
    """Add text to document, handling different markdown elements"""
    lines = text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:], style='Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:], style='Heading 2')
            i += 1
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='Heading 3')
            i += 1
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            doc.add_paragraph('\n'.join(code_lines), style='Normal')
            i += 1

        # Tables
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            if len(table_lines) >= 2:
                # Parse table
                rows = [row.strip().split('|')[1:-1] for row in table_lines]
                rows = [[cell.strip() for cell in row] for row in rows]
                
                # Create table
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                
                # Add header formatting
                for idx, cell in enumerate(table.rows[0].cells):
                    cell.text = rows[0][idx]
                    set_cell_background(cell, 'D3D3D3')
                
                # Add data rows
                for row_idx in range(1, len(rows)):
                    for col_idx in range(len(rows[row_idx])):
                        table.rows[row_idx].cells[col_idx].text = rows[row_idx][col_idx]
        
        # Bullet points
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        
        # Regular paragraphs
        else:
            p = doc.add_paragraph(line)
            i += 1

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to docx"""
    doc = Document()
    md_path = Path(md_file)
    
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split by major sections and process
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:], style='Heading 1')
            i += 1
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:], style='Heading 2')
            i += 1
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='Heading 3')
            i += 1
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            # Add code as a paragraph with monospace formatting
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text, style='Normal')
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1

        # Images
        elif re.match(r'^!\[.*?\]\(.*?\)$', line.strip()):
            match = re.match(r'^!\[(.*?)\]\((.*?)\)$', line.strip())
            if match:
                alt_text, image_path = match.groups()
                resolved_path = (md_path.parent / image_path).resolve()
                if resolved_path.exists():
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(str(resolved_path), width=Inches(6.3))

                    if alt_text:
                        caption = doc.add_paragraph(alt_text)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in caption.runs:
                            run.italic = True
            i += 1
        
        # Tables
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            if len(table_lines) >= 2:
                # Parse table
                rows = [row.strip().split('|')[1:-1] for row in table_lines]
                rows = [[cell.strip() for cell in row] for row in rows]
                
                # Skip separator line
                actual_rows = [rows[0]]
                for row_idx in range(2, len(rows)):
                    actual_rows.append(rows[row_idx])
                
                if actual_rows:
                    # Create table
                    table = doc.add_table(rows=len(actual_rows), cols=len(actual_rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    # Add header formatting
                    for col_idx, cell in enumerate(table.rows[0].cells):
                        if col_idx < len(actual_rows[0]):
                            cell.text = actual_rows[0][col_idx]
                            set_cell_background(cell, 'D3D3D3')
                    
                    # Add data rows
                    for row_idx in range(1, len(actual_rows)):
                        for col_idx in range(len(actual_rows[row_idx])):
                            table.rows[row_idx].cells[col_idx].text = actual_rows[row_idx][col_idx]
        
        # Bullet points
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        
        # Regular paragraphs
        else:
            if line.strip():
                p = doc.add_paragraph(line)
            i += 1
    
    # Save the document
    doc.save(docx_file)
    print(f"✓ Successfully converted {md_file} to {docx_file}")

if __name__ == '__main__':
    md_path = Path(r'c:\Users\thimi\OneDrive - University of Kelaniya\Desktop\real timelog system\docs\REPORT.md')
    docx_path = Path(r'c:\Users\thimi\OneDrive - University of Kelaniya\Desktop\real timelog system\docs\REPORT.docx')
    
    convert_markdown_to_docx(str(md_path), str(docx_path))
